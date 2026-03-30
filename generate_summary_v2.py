#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 创建文档
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 添加标题
title = doc.add_heading('Attention Residuals 技术报告中文总结', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加来源
p = doc.add_paragraph()
run = p.add_run('来源: ')
run.bold = True
p.add_run('Kimi Team (Moonshot AI)')

# 一、研究摘要
doc.add_heading('一、研究摘要', level=1)
doc.add_paragraph(
    '本文提出了注意力残差连接（Attention Residuals, AttnRes），一种替代传统残差连接的新方法。'
    '传统残差连接以固定权重累加各层输出，导致隐藏状态随深度不受控增长，逐层稀释各层的贡献。'
    'AttnRes使用softmax注意力机制对前序层输出进行选择性聚合，允许每层以学习的、依赖输入的权重来聚合早期表示。'
)

# 二、核心问题
doc.add_heading('二、核心问题', level=1)
p = doc.add_paragraph()
run = p.add_run('PreNorm稀释问题：')
run.bold = True
p.add_run('现代LLM普遍采用PreNorm架构配合残差连接，但这种固定单位权重的聚合方式存在以下问题：')

doc.add_paragraph('隐藏状态随深度不受控增长', style='List Bullet')
doc.add_paragraph('各层贡献被逐层稀释', style='List Bullet')
doc.add_paragraph('深层表示的主导性下降', style='List Bullet')

# 三、解决方案
doc.add_heading('三、解决方案', level=1)

doc.add_heading('3.1 完整注意力残差（Full AttnRes）', level=2)
doc.add_paragraph(
    '将固定累加替换为对所有前序层输出的softmax注意力聚合，每层可学习地、依赖输入地选择聚合早期表示。'
    '这解决了稀释问题，但带来内存和通信开销。'
)

doc.add_heading('3.2 分块注意力残差（BlockAttnRes）', level=2)
doc.add_paragraph(
    '将层划分为块，在块级表示上执行注意力，结合基于缓存的流水线通信和两阶段计算策略，'
    '大幅降低内存占用，同时保留大部分性能增益。'
)

# 四、技术细节
doc.add_heading('四、技术细节', level=1)

doc.add_heading('4.1 AttnResOp操作', level=2)
doc.add_paragraph(
    'AttnResOp(α)操作用于聚合前序层输出：h_l = AttnResOp(α)(h_{l-1}, {h_0, ..., h_{l-2}})。'
    '通过注意力权重α实现内容相关的深度选择。'
)

doc.add_heading('4.2 计算优化', level=2)
doc.add_paragraph(
    '采用缓存流水线减少通信开销，两阶段计算策略优化内存使用，使BlockAttnRes成为标准残差连接的实用替代方案。'
)

# 五、实验验证
doc.add_heading('五、实验验证', level=1)

doc.add_heading('5.1 扩展律实验', level=2)
doc.add_paragraph(
    '扩展律实验确认了在不同模型规模下改进的一致性，验证了内容相关深度选择的有效性。'
)

doc.add_heading('5.2 大规模预训练', level=2)
p = doc.add_paragraph('集成到Kimi Linear架构（48B总参数/3B激活参数），在1.4T tokens上预训练：')

doc.add_paragraph('缓解PreNorm稀释问题', style='List Bullet')
doc.add_paragraph('产生更均匀的输出幅度', style='List Bullet')
doc.add_paragraph('更均匀的梯度分布', style='List Bullet')
doc.add_paragraph('下游任务性能全面提升', style='List Bullet')

# 六、主要贡献
doc.add_heading('六、主要贡献', level=1)
doc.add_paragraph('提出AttnRes：用学习式注意力替代固定残差累加', style='List Number')
doc.add_paragraph('提出BlockAttnRes：高效变体，适合大规模训练', style='List Number')
doc.add_paragraph('设计缓存流水线和两阶段计算策略', style='List Number')
doc.add_paragraph('在48B MoE模型上验证有效性', style='List Number')

# 七、结论
doc.add_heading('七、结论', level=1)
doc.add_paragraph(
    'Attention Residuals为现代LLM架构提供了一种即插即用的残差连接替代方案。'
    '通过引入学习式的深度选择机制，AttnRes有效解决了PreNorm架构的稀释问题，'
    '在保持计算效率的同时提升了模型性能。该方法已在Kimi Linear大模型中成功验证，'
    '为未来LLM架构设计提供了新思路。'
)

# 保存文档
output_path = '/Users/javaedge/soft/VSProjects/claude-code-engineering/Attention_Residuals_中文总结_v2.docx'
doc.save(output_path)
print(f'文档已生成: {output_path}')
